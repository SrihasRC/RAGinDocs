import os
import uuid
import base64
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
import asyncio
from datetime import datetime

# Document processing
import PyPDF2
import fitz  # PyMuPDF
import docx
from PIL import Image
import io

# LangChain imports
from langchain.schema.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_groq import ChatGroq

from config.settings import config

class LangChainDocumentProcessor:
    """LangChain-based document processor following multimodal RAG patterns"""
    
    def __init__(self):
        self.temp_dir = config.temp_dir
        
        # Rate limiting configuration to prevent quota exhaustion
        self.max_text_chunks = 10  # Limit text chunks to process with AI
        self.max_tables = 5        # Limit tables to process with AI
        self.max_images = 3        # Limit images to process with AI
        self.use_ai_for_large_docs = False  # Disable AI for docs with too many elements
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.chunk_size,
            chunk_overlap=config.chunk_overlap,
            separators=["\n\n", "\n", " ", ""]
        )
        
        # Initialize LLMs for different tasks
        # Groq for text summarization (fast and higher quota)
        if config.groq_api_key:
            import os
            os.environ["GROQ_API_KEY"] = config.groq_api_key
            self.groq_llm = ChatGroq(
                model="llama-3.1-8b-instant", 
                temperature=0.1
            )
        else:
            self.groq_llm = None
            
        # Gemini for tables and images (better at complex tasks)
        if config.google_api_key:
            self.gemini_llm = ChatGoogleGenerativeAI(
                model="gemini-2.0-flash-exp",
                google_api_key=config.google_api_key,
                temperature=0.1
            )
        else:
            self.gemini_llm = None
            
        # Fallback to whichever is available
        self.llm = self.groq_llm or self.gemini_llm
        
        # Summarization prompt templates
        self.text_summary_prompt = ChatPromptTemplate.from_template("""
        You are an assistant tasked with summarizing document content.
        Give a concise summary that captures the key information and main points.
        
        Content: {content}
        
        Provide only the summary, no additional commentary.
        """)
        
        self.table_summary_prompt = ChatPromptTemplate.from_template("""
        You are an assistant tasked with summarizing tables from documents.
        Describe the table structure, key data points, and insights.
        
        Table content: {content}
        
        Provide only the summary, no additional commentary.
        """)
        
        self.image_summary_prompt = ChatPromptTemplate.from_template("""
        Describe this image in detail. Focus on:
        - Main visual elements and their relationships
        - Any text, charts, or diagrams present
        - The context and purpose of the image
        - Key information that would be useful for answering questions
        
        Image description: {image_info}
        
        Be specific and detailed in your description.
        """)
        
        # Create summarization chains
        if self.llm:
            self.text_chain = self.text_summary_prompt | self.llm | StrOutputParser()
            self.table_chain = self.table_summary_prompt | self.llm | StrOutputParser()
            self.image_chain = self.image_summary_prompt | self.llm | StrOutputParser()
    
    async def process_document(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Process document using LangChain patterns
        
        Returns:
            Dict with processed content and metadata
        """
        doc_id = str(uuid.uuid4())
        file_extension = Path(filename).suffix.lower()
        
        # Save temp file
        temp_file_path = self.temp_dir / f"{doc_id}{file_extension}"
        with open(temp_file_path, "wb") as f:
            f.write(file_content)
        
        try:
            # Extract raw content
            if file_extension == ".pdf":
                raw_content = await self._extract_pdf_content(temp_file_path)
            elif file_extension == ".docx":
                raw_content = await self._extract_docx_content(temp_file_path)
            elif file_extension == ".txt":
                raw_content = await self._extract_txt_content(temp_file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            # Process with LangChain patterns
            processed_content = await self._process_with_langchain(raw_content, doc_id)
            
            # Create metadata
            metadata = {
                "document_id": doc_id,  # Changed from "id" to match vector store expectations
                "filename": filename,   # Changed from "file_name" to match vector store expectations  
                "file_type": file_extension,
                "file_size": len(file_content),
                "total_pages": raw_content.get("page_count"),  # Changed from "page_count"
                "processing_status": "completed",
                "upload_date": str(datetime.now()),
                "processing_date": str(datetime.now()),
                "content_stats": {
                    "text_chunks": len(processed_content["text_documents"]),
                    "tables": len(processed_content["table_documents"]), 
                    "images": len(processed_content["image_documents"])
                }
            }
            
            # Return format expected by vector store
            return {
                "text_documents": processed_content["text_documents"],
                "table_documents": processed_content["table_documents"],
                "image_documents": processed_content["image_documents"],
                "metadata": metadata
            }
            
        except Exception as e:
            raise Exception(f"Document processing failed: {str(e)}")
        finally:
            # Cleanup
            if temp_file_path.exists():
                os.remove(temp_file_path)
    
    async def _extract_pdf_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from PDF using PyMuPDF (better than PyPDF2 for tables/images)"""
        content = {
            "text_chunks": [],
            "tables": [],
            "images": [],
            "page_count": 0
        }
        
        # Use PyMuPDF for comprehensive extraction
        doc = fitz.open(file_path)
        content["page_count"] = len(doc)
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Extract text - handle different PyMuPDF versions gracefully
            text = ""
            try:
                # Standard method for newer versions
                text = page.get_text()  # type: ignore
            except:
                try:
                    # Alternative method
                    text = page.getText()  # type: ignore
                except:
                    # Last resort - skip this page's text
                    print(f"Could not extract text from page {page_num + 1}")
                    continue
                    
            if text.strip():
                content["text_chunks"].append({
                    "page": page_num + 1,
                    "content": text.strip(),
                    "element_id": f"text_{page_num + 1}"
                })
            
            # Extract tables (basic detection)
            try:
                tables = page.find_tables()  # type: ignore
                for table_idx, table in enumerate(tables):
                    table_data = table.extract()  # type: ignore
                    if table_data:
                        content["tables"].append({
                            "page": page_num + 1,
                            "data": table_data,
                            "element_id": f"table_{page_num + 1}_{table_idx}"
                        })
            except Exception as e:
                print(f"Error extracting tables from page {page_num + 1}: {e}")
            
            # Extract images
            try:
                image_list = page.get_images()
                for img_idx, img in enumerate(image_list):
                    xref = img[0]
                    pix = fitz.Pixmap(doc, xref)
                    
                    if pix.n - pix.alpha < 4:  # GRAY or RGB
                        img_data = pix.tobytes("png")
                        img_base64 = base64.b64encode(img_data).decode()
                        
                        content["images"].append({
                            "page": page_num + 1,
                            "base64": img_base64,
                            "element_id": f"image_{page_num + 1}_{img_idx}"
                        })
                    
                    pix = None  # Free memory
            except Exception as e:
                print(f"Error extracting images from page {page_num + 1}: {e}")
        
        doc.close()
        return content
    
    async def _extract_docx_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from DOCX"""
        content = {
            "text_chunks": [],
            "tables": [],
            "images": [],
            "page_count": None
        }
        
        doc = docx.Document(str(file_path))
        
        # Extract paragraphs
        for para_idx, para in enumerate(doc.paragraphs):
            if para.text.strip():
                content["text_chunks"].append({
                    "content": para.text.strip(),
                    "element_id": f"para_{para_idx}"
                })
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            
            if table_data:
                content["tables"].append({
                    "data": table_data,
                    "element_id": f"table_{table_idx}"
                })
        
        # TODO: Extract images from DOCX (requires more complex processing)
        
        return content
    
    async def _extract_txt_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from TXT"""
        content = {
            "text_chunks": [],
            "tables": [],
            "images": [],
            "page_count": None
        }
        
        # Read the text file
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        # Split into paragraphs (chunks separated by double newlines)
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        # If no double newlines, split by single newlines and group
        if len(paragraphs) <= 1 and '\n' in text:
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            # Group lines into chunks of reasonable size
            chunk_size = 5
            paragraphs = []
            for i in range(0, len(lines), chunk_size):
                chunk_lines = lines[i:i+chunk_size]
                paragraphs.append('\n'.join(chunk_lines))
        elif len(paragraphs) <= 1:
            # Single paragraph/line
            paragraphs = [text.strip()] if text.strip() else []
        
        # Add as text chunks
        for para_idx, para in enumerate(paragraphs):
            if para:
                content["text_chunks"].append({
                    "content": para,
                    "element_id": f"para_{para_idx}"
                })
        
        return content
    
    async def _process_with_langchain(self, raw_content: Dict, doc_id: str) -> Dict[str, Any]:
        """Process extracted content using LangChain patterns"""
        
        # Process text chunks
        text_documents = await self._process_text_chunks(raw_content["text_chunks"], doc_id)
        
        # Process tables
        table_documents = await self._process_tables(raw_content["tables"], doc_id)
        
        # Process images
        image_documents = await self._process_images(raw_content["images"], doc_id)
        
        return {
            "text_documents": text_documents,
            "table_documents": table_documents,
            "image_documents": image_documents
        }
    
    async def _process_text_chunks(self, text_chunks: List[Dict], doc_id: str) -> List[Document]:
        """Process text chunks into LangChain Documents with AI rate limiting"""
        documents = []
        
        # Limit the number of chunks to process with AI
        chunks_to_process = text_chunks[:self.max_text_chunks]
        total_chunks = len(text_chunks)
        
        if total_chunks > self.max_text_chunks:
            print(f"⚠️  Rate limit protection: Processing only {self.max_text_chunks} of {total_chunks} text chunks with AI")
        
        for chunk_idx, chunk in enumerate(chunks_to_process):
            # Split large chunks
            chunk_texts = self.text_splitter.split_text(chunk["content"])
            
            for i, text in enumerate(chunk_texts):
                # Generate summary using LangChain (with rate limit protection)
                if chunk_idx < self.max_text_chunks:
                    summary = await self._generate_text_summary(text)
                else:
                    # Use simple truncation for remaining chunks to save API calls
                    summary = text[:500] + "..." if len(text) > 500 else text
                
                # Create Document
                doc = Document(
                    page_content=summary,  # Store summary for retrieval
                    metadata={
                        "doc_id": doc_id,
                        "type": "text",
                        "page": chunk.get("page"),
                        "element_id": f"{chunk.get('element_id', 'text')}_{i}",
                        "original_content": text,  # Store original for context
                        "summary": summary
                    }
                )
                documents.append(doc)
        
        return documents
    
    async def _process_tables(self, tables: List[Dict], doc_id: str) -> List[Document]:
        """Process tables into LangChain Documents with AI rate limiting"""
        documents = []
        
        # Limit the number of tables to process with AI
        tables_to_process = tables[:self.max_tables]
        total_tables = len(tables)
        
        if total_tables > self.max_tables:
            print(f"⚠️  Rate limit protection: Processing only {self.max_tables} of {total_tables} tables with AI")
        
        for table_idx, table in enumerate(tables_to_process):
            # Convert table to text
            table_text = self._table_to_text(table["data"])
            
            # Generate summary with rate limiting
            if table_idx < self.max_tables:
                summary = await self._generate_table_summary(table_text)
            else:
                # Use simple description for remaining tables
                summary = f"Table with {len(table['data'])} rows and {len(table['data'][0]) if table['data'] else 0} columns"
            
            # Create Document
            doc = Document(
                page_content=summary,
                metadata={
                    "doc_id": doc_id,
                    "type": "table",
                    "page": table.get("page"),
                    "element_id": table.get("element_id"),
                    "original_content": table_text,
                    "summary": summary,
                    "table_data": table["data"]  # Keep structured data
                }
            )
            documents.append(doc)
        
        return documents
    
    async def _process_images(self, images: List[Dict], doc_id: str) -> List[Document]:
        """Process images into LangChain Documents with AI rate limiting"""
        documents = []
        
        # Limit the number of images to process with AI
        images_to_process = images[:self.max_images]
        total_images = len(images)
        
        if total_images > self.max_images:
            print(f"⚠️  Rate limit protection: Processing only {self.max_images} of {total_images} images with AI")
        
        for image_idx, image in enumerate(images_to_process):
            # Generate image description with rate limiting
            if image_idx < self.max_images:
                description = await self._generate_image_description(image["base64"])
            else:
                # Use simple description for remaining images
                description = f"Image extracted from document page {image.get('page', 'unknown')}"
            
            # Create Document
            doc = Document(
                page_content=description,
                metadata={
                    "doc_id": doc_id,
                    "type": "image",
                    "page": image.get("page"),
                    "element_id": image.get("element_id"),
                    "image_base64": image["base64"],
                    "description": description
                }
            )
            documents.append(doc)
        
        return documents
    
    async def _generate_text_summary(self, text: str) -> str:
        """Generate text summary using Groq (faster, higher quota) with rate limit protection"""
        try:
            # Use Groq for text summarization if available
            llm_to_use = self.groq_llm or self.gemini_llm
            
            if not llm_to_use:
                return text[:500] + "..." if len(text) > 500 else text
            
            # Create chain with the selected LLM
            text_chain = self.text_summary_prompt | llm_to_use | StrOutputParser()
            result = await text_chain.ainvoke({"content": text})
            return result.strip()
        except Exception as e:
            error_msg = str(e).lower()
            print(f"Error generating text summary: {e}")
            
            # If it's a rate limit error, return a truncated version instead of retrying
            if "quota" in error_msg or "rate limit" in error_msg or "429" in error_msg:
                print("⚠️  Rate limit reached. Using fallback text truncation.")
                return text[:500] + "..." if len(text) > 500 else text
            
            # For other errors, also fallback to truncation
            return text[:500] + "..." if len(text) > 500 else text
    
    async def _generate_table_summary(self, table_text: str) -> str:
        """Generate table summary using Gemini (better at structured data) with rate limit protection"""
        try:
            # Use Gemini for table summarization if available, otherwise Groq
            llm_to_use = self.gemini_llm or self.groq_llm
            
            if not llm_to_use:
                return f"Table with {len(table_text.split())} data points"
            
            # Create chain with the selected LLM
            table_chain = self.table_summary_prompt | llm_to_use | StrOutputParser()
            result = await table_chain.ainvoke({"content": table_text})
            return result.strip()
        except Exception as e:
            error_msg = str(e).lower()
            print(f"Error generating table summary: {e}")
            
            # If it's a rate limit error, return a fallback description
            if "quota" in error_msg or "rate limit" in error_msg or "429" in error_msg:
                print("⚠️  Rate limit reached. Using fallback table description.")
            
            return f"Table with {len(table_text.split())} data points"
    
    async def _generate_image_description(self, image_base64: str) -> str:
        """Generate image description using Gemini with rate limit protection"""
        try:
            # Use Gemini for image description if available, otherwise Groq
            llm_to_use = self.gemini_llm or self.groq_llm
            
            if not llm_to_use:
                return "Image content could not be analyzed"
            
            # For now, we'll create a basic description since we're using text-only models
            # In a full implementation, you'd use a vision model like Gemini Pro Vision
            basic_info = f"Image extracted from document (base64 length: {len(image_base64)})"
            
            # Create chain with the selected LLM
            image_chain = self.image_summary_prompt | llm_to_use | StrOutputParser()
            result = await image_chain.ainvoke({"image_info": basic_info})
            return result.strip()
        except Exception as e:
            error_msg = str(e).lower()
            print(f"Error generating image description: {e}")
            
            # If it's a rate limit error, return a fallback description
            if "quota" in error_msg or "rate limit" in error_msg or "429" in error_msg:
                print("⚠️  Rate limit reached. Using fallback image description.")
            
            return "Image content could not be analyzed"
    
    def _table_to_text(self, table_data: List[List[str]]) -> str:
        """Convert table data to readable text"""
        if not table_data:
            return ""
        
        text_lines = []
        for row in table_data:
            text_lines.append(" | ".join(str(cell) for cell in row))
        
        return "\n".join(text_lines)
